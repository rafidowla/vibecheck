"""
Report generation module for VibeCheck.

Purpose:
    Converts a recording session (transcript + annotated screenshots) into
    structured output files (HTML, DOCX, MD) that can be used as an
    actionable task list for AI coding assistants.

Pipeline:
    1. Organise screenshots into ``img/`` subdirectory.
    2. Call OpenRouter vision API (or fall back to template).
    3. Write ``<slug>.html`` — references ``img/`` with relative paths.
    4. Write ``<slug>.docx`` — screenshots embedded inline within tasks.
    5. Write ``<slug>.md`` — plain Markdown for pasting into an IDE.

Post-generation cleanup (``cleanup_session``):
    - Delete ``recording.wav`` (no longer needed after transcription).
    - Write ``cost.txt`` with AI usage metrics.
    - Rename the session directory from its temp name to a descriptive slug.

Side Effects:
    - HTTP POST to OpenRouter (if API key configured).
    - File I/O: creates ``img/`` directory, writes reports, deletes WAV.

Determinism: Nondeterministic (AI-generated content).
Idempotency: No — each call may produce different AI output.
Thread Safety: Yes — no shared mutable state between calls.
"""

from __future__ import annotations

import base64
import html as html_lib
import io
import logging
import re
import shutil
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Tuple

import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from audit_tool.config import (
    JIRA_CONFIG,
    _get_api_key,
    OPENROUTER_BASE_URL,
    OPENROUTER_MODEL,
    ProcessMode,
)
from audit_tool.transcriber import TranscriptSegment

logger = logging.getLogger(__name__)

# Maximum number of screenshots to send to the vision model.
MAX_SCREENSHOTS_FOR_API: int = 20

# Approximate cost per million tokens for the Qwen VL 72B model.
_COST_PER_1M_INPUT_TOKENS: float = 0.40
_COST_PER_1M_OUTPUT_TOKENS: float = 0.40


@dataclass
class ReportResult:
    """Result of report generation, including cost metrics and Jira keys.

    Attributes:
        report_path: Absolute path to the primary HTML report.
        slug: Descriptive slug used for filenames and folder name.
        model: Model name used for generation (empty if template fallback).
        input_tokens: Number of input tokens consumed.
        output_tokens: Number of output tokens produced.
        cost_usd: Estimated cost in US dollars.
        jira_keys: List of Jira issue keys created (empty if Jira unconfigured).
    """

    report_path: Path
    slug: str = ""
    model: str = ""
    input_tokens: int = 0
    output_tokens: int = 0
    cost_usd: float = 0.0
    jira_keys: list[str] = field(default_factory=list)
    markdown_content: str = ""  # stored so on-demand Jira push doesn't need to re-read disk

    @property
    def cost_display(self) -> str:
        """Human-readable cost string."""
        if self.cost_usd == 0:
            return "Free (template)"
        return f"${self.cost_usd:.4f} ({self.input_tokens:,} in / {self.output_tokens:,} out)"


# -----------------------------------------------------------------------
# Public API
# -----------------------------------------------------------------------


def generate_report(
    session_dir: Path,
    transcript: list[TranscriptSegment],
    clicks: list[ClickRecord],
    mode: ProcessMode = ProcessMode.QA,
) -> ReportResult:
    """Generate a structured feedback document from a recording session.

    Pipeline:
      1. Move all ``click_*.png`` screenshots into ``session_dir/img/``.
      2. Call the AI model with the appropriate mode prompt (or fall back).
      3. Write ``<slug>.html`` with relative ``img/`` references.
      4. Write ``<slug>.docx`` with screenshots inlined within each task/step.
      5. Write ``<slug>.md`` for pasting into an IDE.
      6. If Jira is configured, push tasks/steps as individual issues.

    Args:
        session_dir: Session output directory (may have temp name).
        transcript: Timestamped transcript segments from Whisper.
        clicks: Click records with annotated screenshot paths.
        mode: ``ProcessMode.QA`` (default) → bug/task list for AI agents;
              ``ProcessMode.DOCUMENTATION`` → SOP/tutorial document.

    Returns:
        ``ReportResult`` containing the report path, slug, AI cost, and any
        Jira issue keys created.

    Side Effects:
        - Creates ``img/`` subdirectory and moves PNGs into it.
        - HTTP POST to OpenRouter (if key available).
        - HTTP POSTs to Jira (if JIRA_CONFIG is set).
        - Writes report files to ``session_dir``.
    """
    # ── Step 1: Move screenshots into img/ ──
    _organise_screenshots(session_dir, clicks)

    # ── Step 2: Build the screenshot lookup for HTML ──
    img_lookup = _build_img_lookup(clicks)

    # ── Step 3: Default file paths ──
    docx_path = session_dir / "feedback.docx"
    html_path = session_dir / "feedback.html"
    md_path = session_dir / "feedback.md"
    slug = ""

    # ── Step 4: Generate content via 3-step multi-agent pipeline ──
    if _get_api_key():
        try:
            from audit_tool.pipeline import (
                PipelineError,
                run_documentation_pipeline,
                run_qa_pipeline,
            )

            logger.info("Starting multi-agent pipeline (mode=%s).", mode.value)

            if mode == ProcessMode.DOCUMENTATION:
                markdown_content = run_documentation_pipeline(
                    transcript=transcript,
                    clicks=clicks,
                    session_dir=session_dir,
                    api_key=_get_api_key(),
                    model=OPENROUTER_MODEL,
                )
            else:
                markdown_content = run_qa_pipeline(
                    transcript=transcript,
                    clicks=clicks,
                    session_dir=session_dir,
                    api_key=_get_api_key(),
                    model=OPENROUTER_MODEL,
                )

            # Usage tracking: pipeline makes multiple calls so we report 0
            # for individual token counts (exact tracking is in pipeline logs).
            input_tokens = 0
            output_tokens = 0
            cost_usd = 0.0

            slug = _extract_slug(markdown_content)
            if slug:
                md_path = session_dir / f"{slug}.md"
                html_path = session_dir / f"{slug}.html"
                docx_path = session_dir / f"{slug}.docx"

            md_path.write_text(markdown_content, encoding="utf-8")
            html_content = _wrap_markdown_in_html(markdown_content, img_lookup)
            html_path.write_text(html_content, encoding="utf-8")
            _build_docx_report(docx_path, markdown_content, transcript, clicks)

            # ── Optional Jira push ──
            jira_keys: list[str] = []
            if JIRA_CONFIG is not None:
                try:
                    jira_keys = push_to_jira(
                        JIRA_CONFIG, markdown_content, clicks, mode
                    )
                except Exception as jira_error:
                    logger.error("Jira push failed: %s", jira_error)

            logger.info("Pipeline report saved → %s", html_path)
            return ReportResult(
                report_path=html_path,
                slug=slug,
                model=OPENROUTER_MODEL,
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                cost_usd=cost_usd,
                jira_keys=jira_keys,
                markdown_content=markdown_content,
            )

        except Exception as _pipeline_error:
            import traceback as _tb
            _err_text = _tb.format_exc()
            logger.exception(
                "Pipeline failed (%s) — falling back to legacy single-call path.",
                type(_pipeline_error).__name__,
            )
            # Write error log readable outside the .app bundle
            try:
                _err_log = Path.home() / ".vibecheck" / "error.log"
                _err_log.parent.mkdir(parents=True, exist_ok=True)
                _err_log.write_text(
                    f"Pipeline failed at {__import__('datetime').datetime.now()}:\n{_err_text}\n",
                    encoding="utf-8",
                )
            except Exception:
                pass

            # ── Legacy fallback: original single-call path ──
            try:
                markdown_content, usage = _generate_via_api(
                    transcript, clicks, mode, session_dir
                )
                input_tokens = usage.get("prompt_tokens", 0)
                output_tokens = usage.get("completion_tokens", 0)
                cost_usd = (
                    input_tokens * _COST_PER_1M_INPUT_TOKENS / 1_000_000
                    + output_tokens * _COST_PER_1M_OUTPUT_TOKENS / 1_000_000
                )
                slug = _extract_slug(markdown_content)
                if slug:
                    md_path = session_dir / f"{slug}.md"
                    html_path = session_dir / f"{slug}.html"
                    docx_path = session_dir / f"{slug}.docx"
                md_path.write_text(markdown_content, encoding="utf-8")
                html_content = _wrap_markdown_in_html(markdown_content, img_lookup)
                html_path.write_text(html_content, encoding="utf-8")
                _build_docx_report(docx_path, markdown_content, transcript, clicks)
                jira_keys = []
                if JIRA_CONFIG is not None:
                    try:
                        jira_keys = push_to_jira(JIRA_CONFIG, markdown_content, clicks, mode)
                    except Exception as jira_error:
                        logger.error("Jira push failed: %s", jira_error)
                logger.info("Legacy AI report saved → %s (cost: $%.4f)", html_path, cost_usd)
                return ReportResult(
                    report_path=html_path,
                    slug=slug,
                    model=OPENROUTER_MODEL,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    cost_usd=cost_usd,
                    jira_keys=jira_keys,
                    markdown_content=markdown_content,
                )
            except Exception:
                logger.exception("Legacy single-call path also failed — falling back to template.")

    # ── Fallback: template-based report ──
    markdown_content = _generate_template_report(transcript, clicks, mode)
    slug = _extract_slug_from_transcript(transcript)
    if slug:
        md_path = session_dir / f"{slug}.md"
        html_path = session_dir / f"{slug}.html"
        docx_path = session_dir / f"{slug}.docx"

    md_path.write_text(markdown_content, encoding="utf-8")
    html_content = _build_template_html(transcript, clicks, img_lookup)
    html_path.write_text(html_content, encoding="utf-8")
    _build_docx_report(docx_path, markdown_content, transcript, clicks)

    logger.info("Template report saved → %s", html_path)
    return ReportResult(report_path=html_path, slug=slug, markdown_content=markdown_content)


def cleanup_session(session_dir: Path, result: ReportResult) -> Path:
    """Finalise a session directory after report generation.

    1. Deletes ``recording.wav`` (no longer needed after transcription).
    2. Writes ``cost.txt`` with AI usage metrics.
    3. Renames the directory to the AI-generated slug.

    Args:
        session_dir: The session output directory (still has temp name).
        result: The ``ReportResult`` from ``generate_report()``.

    Returns:
        The final session directory path (may differ if renamed).

    Side Effects:
        Deletes WAV, writes cost.txt, renames directory.

    Determinism: Deterministic.
    Idempotency: Safe to call multiple times.
    Thread Safety: Yes.
    """
    # ── Delete recording.wav ──
    wav_file = session_dir / "recording.wav"
    if wav_file.exists():
        try:
            wav_file.unlink()
            logger.info("Deleted raw audio: %s", wav_file)
        except OSError as err:
            logger.warning("Could not delete WAV: %s", err)

    # ── Write cost.txt ──
    cost_path = session_dir / "cost.txt"
    cost_lines = [
        f"Model:         {result.model or 'Template (no AI)'}",
        f"Input tokens:  {result.input_tokens:,}",
        f"Output tokens: {result.output_tokens:,}",
        f"Est. cost:     {result.cost_display}",
    ]
    cost_path.write_text("\n".join(cost_lines) + "\n", encoding="utf-8")
    logger.info("Cost summary → %s", cost_path)

    # ── Rename directory to descriptive slug with date suffix ──
    final_dir = session_dir
    # Append YYYYMMDD as suffix for human readability (name visible first when truncated):
    #   e.g. "login-page-20260319" instead of "20260319-login-page"
    date_suffix = datetime.now().strftime("%Y%m%d")
    if result.slug:
        dated_slug = f"{result.slug}-{date_suffix}"
        final_dir = _rename_session_dir(session_dir, dated_slug)
    else:
        # Fallback: clean up the temp name to a readable HHMMSS-YYYYMMDD format
        fallback_slug = datetime.now().strftime("%H%M%S-%Y%m%d")
        final_dir = _rename_session_dir(session_dir, fallback_slug)

    return final_dir


# -----------------------------------------------------------------------
# Screenshot organisation
# -----------------------------------------------------------------------


def _organise_screenshots(session_dir: Path, clicks: list) -> None:
    """Move all click_*.png files into an img/ subdirectory.

    Updates each ClickRecord.screenshot_path in place.

    Args:
        session_dir: The session output directory.
        clicks: Click records whose paths will be updated.

    Side Effects:
        Creates img/ directory and moves PNGs.
    """
    img_dir = session_dir / "img"
    img_dir.mkdir(exist_ok=True)
    moved = 0
    for click in clicks:
        src = click.screenshot_path
        if src.exists() and src.parent != img_dir:
            dst = img_dir / src.name
            try:
                shutil.move(str(src), str(dst))
                click.screenshot_path = dst
                moved += 1
            except OSError as err:
                logger.warning("Could not move %s: %s", src.name, err)
    logger.info("Moved %d screenshots → %s", moved, img_dir)


def _build_img_lookup(clicks: list) -> dict[str, str]:
    """Build a filename → relative path lookup for HTML references.

    Returns:
        Dict mapping e.g. 'click_0001.png' → 'img/click_0001.png'.
    """
    lookup: dict[str, str] = {}
    for click in clicks:
        if click.screenshot_path.exists():
            lookup[click.screenshot_path.name] = f"img/{click.screenshot_path.name}"
    return lookup


# -----------------------------------------------------------------------
# DOCX generation — screenshots inlined within tasks
# -----------------------------------------------------------------------


def _build_docx_report(
    docx_path: Path,
    markdown_text: str,
    transcript: list[TranscriptSegment],
    clicks: list,
) -> None:
    """Build a Word document with screenshots inlined within each task.

    Structure:
      1. Title + timestamp
      2. AI-generated task list with screenshots embedded inline
         (when a task references click_NNNN.png, the image appears
         right there in the task).
      3. Transcript section at the end.

    Args:
        docx_path: Output path for the .docx file.
        markdown_text: The Markdown report content.
        transcript: Timestamped transcript segments.
        clicks: Click records with screenshot paths (in img/).

    Side Effects:
        Writes the DOCX file to disk.
    """
    doc = Document()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Build a lookup: filename → absolute path
    img_paths: dict[str, Path] = {}
    for click in clicks:
        if click.screenshot_path.exists():
            img_paths[click.screenshot_path.name] = click.screenshot_path

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # ── Title ──
    title_para = doc.add_heading("VibeCheck — Tasks", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Generated: {now}", style="Subtitle")
    doc.add_paragraph("")

    # ── Task list with inline screenshots ──
    _add_markdown_to_docx(doc, markdown_text, img_paths)

    # ── Transcript appendix ──
    doc.add_page_break()
    doc.add_heading("Appendix: Transcript", level=1)
    if transcript:
        for seg in transcript:
            start_ts = _seconds_to_mmss(seg.start)
            end_ts = _seconds_to_mmss(seg.end)
            para = doc.add_paragraph()
            ts_run = para.add_run(f"[{start_ts} → {end_ts}]  ")
            ts_run.font.name = "Courier New"
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(0x33, 0x99, 0x77)
            text_run = para.add_run(seg.text)
            text_run.font.size = Pt(10)
    else:
        doc.add_paragraph("No speech was recorded.", style="Intense Quote")

    doc.save(str(docx_path))
    logger.info("DOCX report saved → %s", docx_path)


def _add_markdown_to_docx(
    doc: Document,
    markdown_text: str,
    img_paths: dict[str, Path],
) -> None:
    """Convert Markdown into DOCX paragraphs with inline screenshots.

    When a line references a screenshot filename (click_NNNN.png), the
    image is embedded directly below that line in the document.

    Args:
        doc: The python-docx Document to append to.
        markdown_text: Raw Markdown string.
        img_paths: Mapping of screenshot filename → absolute Path.

    Side Effects:
        Appends paragraphs and images to the Document.
    """
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Check if this line references a screenshot
        referenced_img = _extract_img_reference(stripped, img_paths)

        # Headings
        if stripped.startswith("# "):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            level = min(hashes, 4)
            text = stripped.lstrip("# ").strip()
            doc.add_heading(text, level=level)

        # Checkbox items
        elif stripped.startswith("- [ ] "):
            text = stripped[6:].strip()
            if text:
                para = doc.add_paragraph(style="List Bullet")
                para.add_run("☐ ").font.size = Pt(12)
                para.add_run(text)

        elif stripped.startswith("- [x] "):
            text = stripped[6:].strip()
            if text:
                para = doc.add_paragraph(style="List Bullet")
                check_run = para.add_run("☑ ")
                check_run.font.size = Pt(12)
                check_run.font.color.rgb = RGBColor(0x33, 0x99, 0x77)
                para.add_run(text)

        # Bold key-value lines (- **Priority:** High)
        elif stripped.startswith("- **"):
            text = stripped[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            # Parse **Key:** Value
            bold_match = re.match(r"\*\*(.+?)\*\*\s*(.*)", text)
            if bold_match:
                key_run = para.add_run(bold_match.group(1) + " ")
                key_run.bold = True
                key_run.font.size = Pt(10)
                val_run = para.add_run(bold_match.group(2))
                val_run.font.size = Pt(10)
            else:
                para.add_run(text.replace("**", ""))

        # Regular bullet items
        elif stripped.startswith("- "):
            text = stripped[2:].strip()
            if text:
                doc.add_paragraph(text, style="List Bullet")

        # Blockquotes
        elif stripped.startswith("> "):
            text = stripped[2:].strip()
            if text:
                doc.add_paragraph(text, style="Intense Quote")

        # Plain text
        else:
            doc.add_paragraph(stripped)

        # ── Inline screenshot if this line referenced one ──
        if referenced_img and referenced_img.exists():
            doc.add_picture(str(referenced_img), width=Inches(5.5))
            doc.add_paragraph("")  # spacer


def _extract_img_reference(
    line: str,
    img_paths: dict[str, Path],
) -> Path | None:
    """Check if a line references a screenshot filename.

    Args:
        line: A single line of Markdown text.
        img_paths: Mapping of screenshot filename → Path.

    Returns:
        The Path to the referenced image, or None.
    """
    for filename in img_paths:
        if filename in line:
            return img_paths[filename]
    return None


# -----------------------------------------------------------------------
# AI-powered generation via OpenRouter
# -----------------------------------------------------------------------


def _generate_via_api(
    transcript: list[TranscriptSegment],
    clicks: list,
    mode: ProcessMode = ProcessMode.QA,
    session_dir: Optional[Path] = None,
) -> Tuple[str, dict]:
    """Call the OpenRouter chat completions endpoint with vision.

    Sends the prompt instructions as a ``system`` message and the session
    data (transcript, click log, screenshots) as the ``user`` message.
    This separation prevents models from treating the instructions as
    content to echo back.

    Args:
        transcript: Timestamped speech segments.
        clicks: Click records with screenshot paths.
        mode: Process mode that selects the system prompt.

    Returns:
        A tuple of (markdown_content, usage_dict).

    Raises:
        httpx.HTTPStatusError: On non-2xx responses.
        httpx.TimeoutException: On network timeout.
        RuntimeError: If the AI response fails validation.
    """
    transcript_text = _format_transcript(transcript)

    if mode == ProcessMode.DOCUMENTATION:
        system_text, user_text = _build_documentation_prompt(
            transcript_text, transcript, clicks, session_dir
        )
    else:
        system_text, user_text = _build_qa_prompt(
            transcript_text, transcript, clicks, session_dir
        )

    # Build the user content array: text + screenshot images
    user_content: list[dict[str, Any]] = [
        {
            "type": "text",
            "text": user_text,
        }
    ]

    screenshots_sent = 0
    for click in clicks:
        if screenshots_sent >= MAX_SCREENSHOTS_FOR_API:
            break
        img_path = click.screenshot_path
        if not img_path.exists():
            logger.warning("Screenshot missing, skipping: %s", img_path)
            continue
        b64 = base64.b64encode(img_path.read_bytes()).decode("ascii")
        # Label the image so the model can correctly reference the filename.
        # Without this, the model guesses which click_XXXX.png maps to which image.
        user_content.append(
            {
                "type": "text",
                "text": f"[Image: {img_path.name} — Click #{click.index}]",
            }
        )
        user_content.append(
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{b64}",
                },
            }
        )
        screenshots_sent += 1

    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [
            {"role": "system", "content": system_text},
            {"role": "user", "content": user_content},
        ],
        "max_tokens": 8192,
        "temperature": 0.3,
    }

    headers = {
        "Authorization": f"Bearer {_get_api_key()}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/vibecheck",
        "X-Title": "VibeCheck",
    }

    response = httpx.post(
        OPENROUTER_BASE_URL,
        json=payload,
        headers=headers,
        timeout=120.0,
    )
    response.raise_for_status()
    data = response.json()

    markdown_text = data["choices"][0]["message"]["content"]
    usage = data.get("usage", {})
    logger.info(
        "API usage — input: %d tokens, output: %d tokens",
        usage.get("prompt_tokens", 0),
        usage.get("completion_tokens", 0),
    )

    # ── Validate AI response ──
    # If the response doesn't contain any heading or structural markers,
    # the model likely ignored the prompt and dumped raw content.
    if mode == ProcessMode.QA:
        expected_markers = ["Task", "#"]
    else:
        expected_markers = ["Step", "#"]

    has_structure = any(marker in markdown_text for marker in expected_markers)
    if not has_structure:
        logger.warning(
            "AI response lacks expected structure markers %s — "
            "response may not follow the prompt. Falling back to template.",
            expected_markers,
        )
        raise RuntimeError("AI response did not follow prompt structure.")

    return markdown_text, usage


from audit_tool.config import resource_path as _resource_path

_PROMPTS_DIR = _resource_path("prompts")


def _load_prompt(
    filename: str,
    transcript_text: str,
    click_summary: str,
) -> Tuple[str, str]:
    """Load a prompt template from the ``prompts/`` directory and substitute values.

    Purpose:
        Allows users to customise the AI prompt without editing Python source.
        Reads ``prompts/<filename>`` and substitutes ``{{TRANSCRIPT}}`` and
        ``{{CLICKS}}`` placeholders.

        If the file contains a ``---`` separator line, everything above it is
        treated as the **system prompt** (instructions) and everything below
        as the **user data**.  If no separator exists, the entire content is
        used as the system prompt and the user data is left empty.

    Args:
        filename: Prompt file name, e.g. ``"qa_prompt.md"``.
        transcript_text: Formatted transcript to inject.
        click_summary: Formatted click position summary to inject.

    Returns:
        A ``(system_text, user_text)`` tuple with placeholders substituted.
        Returns ``("", "")`` if the file cannot be read (caller uses default).

    Side Effects:
        Reads one file from disk.
    """
    prompt_path = _PROMPTS_DIR / filename
    try:
        template = prompt_path.read_text(encoding="utf-8")
        rendered = (
            template
            .replace("{{TRANSCRIPT}}", transcript_text)
            .replace("{{CLICKS}}", click_summary)
        )
        logger.info("Loaded custom prompt from %s", prompt_path)
        return rendered, ""
    except OSError:
        logger.debug("Prompt file not found: %s — using built-in default.", prompt_path)
        return "", ""


def _build_correlated_click_log(
    transcript: list,
    clicks: list,
    session_dir: Optional[Path] = None,
) -> str:
    """Build a click log where each click is annotated with concurrent speech.

    Purpose:
        Closes the time-axis gap between TranscriptSegment.start (seconds from
        recording start) and ClickRecord.timestamp (Unix epoch).  The result
        is an annotated click log the model can use to pick the most
        contextually relevant screenshot for each issue described in the narration.

    Algorithm:
        1. Read recording_start.txt (written by AudioRecorder.start()) for the
           true recording start epoch.  This correctly handles sessions where
           the user narrated for several seconds before clicking.
           Falls back to min(click_epoch) - earliest_transcript_start if the
           file is absent (backwards compatibility with old sessions).
        2. For each click, convert its epoch to session-relative seconds.
        3. Find all transcript segments whose window overlaps ±5 s around the
           click.  A wider window handles speech that started before/after a click.
        4. Format: "Click #N @ T+X.Xs → \"spoken text\" — file=click_NNNN.png"

    Args:
        transcript: Ordered list of TranscriptSegment with .start and .end.
        clicks: Click records with .timestamp (Unix epoch) and .screenshot_path.
        session_dir: Session directory containing recording_start.txt.

    Returns:
        A multi-line string describing each click with its concurrent speech context.

    Side Effects: None.
    Determinism: Deterministic.
    Thread Safety: Yes (read-only).
    """
    if not clicks:
        return "  (no clicks recorded)"

    # ── Derive session_start_epoch ──────────────────────────────────────
    # Primary: read the exact epoch that AudioRecorder.start() stamped.
    # Fallback: estimate from min(click epoch) - earliest transcript start
    # (this under-anchors if the user narrated before clicking, but beats nothing).
    session_start_epoch: float = 0.0
    anchor_source = "fallback"

    if session_dir is not None:
        start_file = session_dir / "recording_start.txt"
        if start_file.exists():
            try:
                session_start_epoch = float(start_file.read_text(encoding="utf-8").strip())
                anchor_source = "recording_start.txt"
            except Exception:
                pass

    if anchor_source == "fallback":
        earliest_click: float = min(c.timestamp for c in clicks)
        earliest_transcript: float = min((seg.start for seg in transcript), default=0.0)
        session_start_epoch = earliest_click - earliest_transcript

    lines: list[str] = [
        f"  [anchor={anchor_source}, session_start_epoch={session_start_epoch:.3f}]"
    ]
    for click in clicks:
        click_session_t = click.timestamp - session_start_epoch

        # Find transcript segments concurrent with (or within 5 s of) this click.
        # A wider window handles the natural delay between speaking and clicking.
        window = 5.0
        concurrent: list[str] = [
            seg.text.strip()
            for seg in transcript
            if seg.start - window <= click_session_t <= seg.end + window
        ]
        speech = (
            '"' + " ".join(concurrent) + '"'
            if concurrent
            else "(between speech segments)"
        )

        lines.append(
            f"  - Click #{click.index} @ T+{click_session_t:.1f}s"
            f" → {speech}"
            f" — file={click.screenshot_path.name}"
            f", pos=({click.x:.0f}, {click.y:.0f})"
        )



    return "\n".join(lines)


def _build_qa_prompt(
    transcript_text: str,
    transcript: list,
    clicks: list,
    session_dir: Optional[Path] = None,
) -> tuple:
    """Return the QA mode prompt as a (system, user) pair.

    Purpose:
        Generates a structured bug/task list for AI coding agents.  The
        temporally-correlated click log lets the model pick the screenshot
        closest to when each issue was spoken — not just the nearest by index.

    Args:
        transcript_text: Pre-formatted timestamped transcript string.
        transcript: Raw TranscriptSegment list for temporal correlation.
        clicks: Click records with screenshot paths.
        session_dir: Session directory containing recording_start.txt.

    Returns:
        A (system_text, user_text) tuple.
    """
    click_summary = _build_correlated_click_log(transcript, clicks, session_dir)

    # Try file-based prompt first
    custom_system, custom_user = _load_prompt("qa_prompt.md", transcript_text, click_summary)
    if custom_system:
        return custom_system, custom_user

    # Built-in fallback — system prompt (instructions)
    system_text = """You are a senior software engineer performing a structured QA/QC review.

I recorded a screen review session narrating issues while clicking through the application.
You have my narration, a temporally-correlated click log, and the annotated screenshots.

## Your Role

**All tasks must be based solely on what I SPOKE.** Do not infer new bugs or fixes from the screenshots.
Use screenshots only to:
- Select the single best screenshot for each task using the concurrent speech context
- Deduplicate jittery or repeated clicks on the same area — choose one representative screenshot

## Screenshot Selection Rules
1. Each click in the log is annotated with what was being spoken at that moment.
2. For each task, pick the screenshot whose spoken-text annotation best matches the issue described.
3. If multiple clicks share similar speech context, pick the one closest in time to the issue.
4. If no click is contextually close to a spoken issue, write `(no screenshot)` for that task.

## Output Format

Produce a Markdown document with this EXACT structure:

```
# [App/Feature Name] — QA Tasks

## Summary
2-3 sentences: what area was reviewed, biggest issues found from the narration.

## Tasks

### Task 1: [Clear, specific title from narration]
- **Priority:** Critical / High / Medium / Low
- **Type:** Bug | UI | UX | Missing Feature | Performance
- **Screenshot:** click_NNNN.png
- **Target Component:** `[Best guess at file/component, e.g., Sidebar.tsx or header.css]`
- **What's wrong:** [Exactly what I described verbally.]
- **Implementation steps:**
  1. Open `[likely filename or component]`
  2. Locate the [element] responsible for [behavior]
  3. Change [specific property] from [current] to [target]
- **Acceptance criteria:**
  - [ ] [Specific, testable condition]

### Task 2: …
```

## Critical Rules
1. All tasks must come from the narration — not from reading screenshots.
2. Each task must be independently actionable by an AI coding agent.
3. Implementation steps must name specific files, components, props, and values.
4. Do not guess exact hex codes, pixel values, or font sizes unless I stated them.
5. Consolidate related micro-issues affecting the same component into a single task.
6. Acceptance criteria must be specific and testable.
7. Number tasks sequentially. Output ONLY the Markdown document."""

    # User message — session data
    user_text = f"""## My Spoken Observations (timestamped)
{transcript_text}

## Click Log — each click annotated with concurrent speech
(Use the spoken text next to each click to pick the best screenshot for each task.
 The click whose speech annotation matches the issue is the right screenshot to use.)
{click_summary}"""

    return system_text, user_text




def _build_documentation_prompt(
    transcript_text: str,
    transcript: list,
    clicks: list,
    session_dir: Optional[Path] = None,
) -> tuple:
    """Return the Documentation mode prompt as a (system, user) pair.

    Purpose:
        Generates an instructional SOP / tutorial document.  Voice narration is
        the primary content source.  The correlated click log maps each screenshot
        to the speech spoken at that moment so the model pairs steps with the
        most contextually relevant screenshot.

    Args:
        transcript_text: Pre-formatted timestamped transcript.
        transcript: Raw TranscriptSegment list for temporal correlation.
        clicks: Click records (for step-by-step pairing).
        session_dir: Session directory containing recording_start.txt.

    Returns:
        A (system_text, user_text) tuple.
    """
    click_summary = _build_correlated_click_log(transcript, clicks, session_dir)

    # Try file-based prompt first
    custom_system, custom_user = _load_prompt("documentation_prompt.md", transcript_text, click_summary)
    if custom_system:
        return custom_system, custom_user

    # Built-in fallback — system prompt (instructions)
    system_text = """You are a senior technical writer converting a recorded walkthrough into a how-to guide.

You have the speaker's narration and annotated screenshots of each step.
Each screenshot in the click log is annotated with what was being spoken at the moment of the click.

## Your Role

**Voice narration dictates the core steps.** Every step in the guide must correlate to an action that was spoken.
However, you MUST use the visual evidence in the screenshots to drastically enrich the descriptions of the steps. Make your documentation far more descriptive, accurate, and professional than my rough verbal narration by observing the exact UI state shown in the screenshots.

Use the screenshots to:
- Drastically improve the descriptiveness of the UI elements, using exact labels, exact icon shapes, button colors, and visual layout details visible in the image.
- Fill in precise UI element names where narration was vague (e.g. "I clicked there" → "Click the blue **Save as Draft** button next to the title").
- Match each step to the screenshot whose speech annotation best matches the narrated action.

Do NOT:
- Add extra functional steps or entirely new concepts not mentioned or implied by the narration.
- Hallucinate UI elements or features that are not clearly visible in the screenshots.

## Output Format

Produce a Markdown how-to guide:
- Overview (1 sentence)
- Prerequisites (from narration only, or "None")
- Steps (### Step N with screenshot reference and 1–3 sentence description in second person)
- Tips & Gotchas (only if explicitly mentioned)
- ✅ Done when (single testable condition)

Write in second person. Output ONLY the Markdown document."""

    # User message — session data
    user_text = f"""## Narration (timestamped) — PRIMARY SOURCE
{transcript_text}

## Screenshot Sequence — each annotated with concurrent speech
(Use the spoken text next to each screenshot to pair steps with the correct image.)
{click_summary}"""

    return system_text, user_text



_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>VibeCheck — {title}</title>
<style>
  :root {{ --bg: #0f0f14; --surface: #1a1a24; --border: #2a2a3a; --fg: #e0e0e0; --fg2: #999; --accent: #4ecca3; --red: #e94560; }}
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ background:var(--bg); color:var(--fg); font-family:'Inter','Helvetica Neue',sans-serif; line-height:1.7; padding:2rem; max-width:960px; margin:0 auto; }}
  h1 {{ color:var(--accent); font-size:1.8rem; margin-bottom:0.5rem; border-bottom:2px solid var(--border); padding-bottom:0.5rem; }}
  h2 {{ color:var(--fg); font-size:1.3rem; margin-top:2rem; margin-bottom:0.5rem; }}
  h3 {{ color:var(--fg2); font-size:1.1rem; margin-top:1.5rem; margin-bottom:0.3rem; }}
  p, li {{ color:var(--fg); margin-bottom:0.5rem; }}
  ul {{ padding-left:1.5rem; }}
  code {{ background:var(--surface); padding:2px 6px; border-radius:4px; font-size:0.9rem; color:var(--accent); }}
  strong {{ color: #f0f0f0; }}
  .screenshot {{ margin:1rem 0; background:var(--surface); border:1px solid var(--border); border-radius:8px; overflow:hidden; }}
  .screenshot img {{ width:100%; display:block; }}
  .screenshot .caption {{ padding:0.5rem 1rem; color:var(--fg2); font-size:0.85rem; border-top:1px solid var(--border); }}
  .task-item {{ display:flex; align-items:flex-start; gap:0.5rem; margin:0.3rem 0; }}
  .task-item input[type=checkbox] {{ margin-top:0.35rem; accent-color:var(--accent); }}
  blockquote {{ border-left:3px solid var(--accent); padding-left:1rem; color:var(--fg2); margin:1rem 0; }}
  .transcript {{ background:var(--surface); border:1px solid var(--border); border-radius:8px; padding:1rem; margin:1rem 0; font-size:0.9rem; max-height:400px; overflow-y:auto; }}
  .transcript .seg {{ margin-bottom:0.4rem; }}
  .transcript .ts {{ color:var(--accent); font-family:monospace; font-size:0.8rem; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def _wrap_markdown_in_html(
    markdown_text: str,
    img_lookup: dict[str, str],
) -> str:
    """Wrap AI-generated Markdown into HTML with inline screenshots.

    When the Markdown references ``click_NNNN.png``, it is replaced with
    an ``<img>`` tag pointing to the relative ``img/`` path.

    Args:
        markdown_text: The Markdown content from the AI model.
        img_lookup: Mapping of filename → relative path (e.g. 'img/click_0001.png').

    Returns:
        A complete HTML string.
    """
    body_html = _markdown_to_simple_html(markdown_text)

    # Replace screenshot filename references with inline images
    for filename, rel_path in img_lookup.items():
        img_html = (
            f'</p><div class="screenshot">'
            f'<img src="{rel_path}" alt="{filename}">'
            f'<div class="caption">📸 {filename}</div>'
            f'</div><p>'
        )
        body_html = body_html.replace(filename, img_html)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    return _HTML_TEMPLATE.format(title=now, body=body_html)


def _markdown_to_simple_html(md_text: str) -> str:
    """Convert basic Markdown to HTML (lightweight, no dependency).

    Handles headings, bullet lists, checkboxes, bold, code, blockquotes.
    """
    lines = md_text.split("\n")
    html_lines: list[str] = []
    in_list = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("# ").strip()
            html_lines.append(f"<h{level}>{html_lib.escape(text)}</h{level}>")

        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            checked = "checked" if stripped.startswith("- [x]") else ""
            text = stripped[6:].strip()
            html_lines.append(
                f'<li class="task-item"><input type="checkbox" {checked}> {html_lib.escape(text)}</li>'
            )

        elif stripped.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            text = stripped[2:].strip()
            # Preserve bold markers
            text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html_lib.escape(text))
            html_lines.append(f"<li>{text}</li>")

        elif stripped.startswith("> "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            text = stripped[2:].strip()
            html_lines.append(f"<blockquote>{html_lib.escape(text)}</blockquote>")

        elif not stripped:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<br>")

        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html_lib.escape(stripped))
            html_lines.append(f"<p>{text}</p>")

    if in_list:
        html_lines.append("</ul>")

    return "\n".join(html_lines)


# -----------------------------------------------------------------------
# Fallback: template-based reports
# -----------------------------------------------------------------------


def _generate_template_report(
    transcript: list[TranscriptSegment],
    clicks: list,
    mode: ProcessMode = ProcessMode.QA,
) -> str:
    """Build a structured Markdown report without an LLM.

    Args:
        transcript: Timestamped speech segments.
        clicks: Click records.
        mode: Process mode, used to tailor the template header and task list
            labels (QA vs Documentation).

    Returns:
        A Markdown string ready to write to disk.
    """
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    transcript_text = _format_transcript(transcript)

    click_section = ""
    for click in clicks:
        click_section += (
            f"### Click #{click.index} — ({click.x}, {click.y})\n"
            f"- **Time:** {_epoch_to_time(click.timestamp)}\n"
            f"- **Screenshot:** `{click.screenshot_path.name}`\n\n"
        )

    if mode == ProcessMode.DOCUMENTATION:
        header = f"# Documentation Session — {now}"
        task_label = "## Steps (fill in manually)"
        task_placeholder = "- [ ] Step 1: "
    else:
        header = f"# QA Review — Session {now}"
        task_label = "## Task List (fill in manually)"
        task_placeholder = "- [ ] "

    return f"""{header}

> **Note:** This report was generated offline (no AI API key configured).
> Review the transcript and screenshots below to complete this document.

## Transcript

{transcript_text if transcript_text.strip() else "_No speech was recorded._"}

## Click / Screenshot Log

{click_section if click_section.strip() else "_No clicks were recorded._"}

{task_label}

{task_placeholder}
"""


def _build_template_html(
    transcript: list[TranscriptSegment],
    clicks: list,
    img_lookup: dict[str, str],
) -> str:
    """Build a template HTML report (no LLM) with img/ references."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    if transcript:
        transcript_html = '<div class="transcript">'
        for seg in transcript:
            start_ts = _seconds_to_mmss(seg.start)
            end_ts = _seconds_to_mmss(seg.end)
            transcript_html += (
                f'<div class="seg">'
                f'<span class="ts">[{start_ts} → {end_ts}]</span> '
                f'{html_lib.escape(seg.text)}'
                f'</div>'
            )
        transcript_html += '</div>'
    else:
        transcript_html = '<p><em>No speech was recorded.</em></p>'

    screenshots_html = ""
    for click in clicks:
        fname = click.screenshot_path.name
        rel_path = img_lookup.get(fname, f"img/{fname}")
        screenshots_html += (
            f'<div class="screenshot">'
            f'<img src="{rel_path}" alt="{fname}">'
            f'<div class="caption">'
            f'📸 Click #{click.index} at ({click.x}, {click.y}) — {_epoch_to_time(click.timestamp)}'
            f'</div></div>'
        )

    if not screenshots_html:
        screenshots_html = '<p><em>No clicks were recorded.</em></p>'

    body = f"""
<h1>🎯 VibeCheck — Session {now}</h1>
<blockquote>Review the transcript and screenshots below. Fill in the task list as needed.</blockquote>

<h2>📝 Transcript</h2>
{transcript_html}

<h2>📸 Annotated Screenshots</h2>
{screenshots_html}

<h2>✅ Task List</h2>
<ul><li class="task-item"><input type="checkbox"> </li></ul>
"""

    return _HTML_TEMPLATE.format(title=now, body=body)


# -----------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------


def _format_transcript(segments: list[TranscriptSegment]) -> str:
    """Format transcript segments into a readable timestamped block."""
    if not segments:
        return "_No transcript available._"
    lines: list[str] = []
    for seg in segments:
        start_ts = _seconds_to_mmss(seg.start)
        end_ts = _seconds_to_mmss(seg.end)
        lines.append(f"[{start_ts} → {end_ts}] {seg.text}")
    return "\n".join(lines)


def _seconds_to_mmss(seconds: float) -> str:
    """Convert seconds to MM:SS format."""
    minutes = int(seconds) // 60
    secs = int(seconds) % 60
    return f"{minutes:02d}:{secs:02d}"


def _epoch_to_time(epoch: float) -> str:
    """Convert a Unix timestamp to HH:MM:SS."""
    return datetime.fromtimestamp(epoch).strftime("%H:%M:%S")


def _slugify(text: str, max_length: int = 50) -> str:
    """Convert text into a filename-safe slug."""
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    text = re.sub(r"-+", "-", text)
    return text[:max_length].rstrip("-")


def _extract_slug(markdown_text: str) -> str:
    """Extract a descriptive slug from the AI-generated Markdown title.

    Handles both QA mode (e.g. ``# App — QA Tasks``) and Documentation
    mode (e.g. ``# Feature — How-To Guide``) title suffixes.
    """
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped.lstrip("# ").strip()
            for suffix in [
                "— QA Tasks", "- QA Tasks", ": QA Tasks",
                "— Audit Tasks", "- Audit Tasks", ": Audit Tasks",
                "— How-To Guide", "- How-To Guide", ": How-To Guide",
                "— Tutorial", "- Tutorial", ": Tutorial",
                "— SOP", "- SOP", ": SOP",
                "— Documentation", "- Documentation",
            ]:
                if title.lower().endswith(suffix.lower()):
                    title = title[: -len(suffix)].strip()
                    break
            for prefix in [
                "Audit Feedback —", "Audit Feedback -", "Audit Feedback:",
                "Audit Tasks —", "Audit Tasks -",
                "QA Review —", "QA Review -",
                "VibeCheck —", "VibeCheck -",
                "How-To Guide —", "How-To Guide -",
                "Tutorial —", "Tutorial -",
            ]:
                if title.lower().startswith(prefix.lower()):
                    title = title[len(prefix):].strip()
                    break
            slug = _slugify(title)
            if slug and slug not in (
                "audit-feedback", "audit-tasks", "qa-review",
                "vibecheck", "how-to-guide", "tutorial",
            ):
                return slug
    return ""


def _extract_slug_from_transcript(transcript: list[TranscriptSegment]) -> str:
    """Derive a rough slug from the first few transcript words."""
    if not transcript:
        return ""
    all_text = " ".join(seg.text for seg in transcript[:3])
    words = all_text.split()[:8]
    if not words:
        return ""
    return _slugify(" ".join(words))


def _rename_session_dir(session_dir: Path, slug: str) -> Path:
    """Rename the session directory to a clean, descriptive name.

    Replaces the temporary ``_recording_<timestamp>`` directory with a
    slug.  Appends ``-2``, ``-3``, etc. on collisions.

    Args:
        session_dir: Current session directory path.
        slug: The descriptive slug to use as the folder name.

    Returns:
        The new directory path.

    Side Effects:
        Renames the directory on disk.
    """
    parent = session_dir.parent
    new_path = parent / slug

    if new_path.exists():
        counter = 2
        while True:
            candidate = parent / f"{slug}-{counter}"
            if not candidate.exists():
                new_path = candidate
                break
            counter += 1

    try:
        session_dir.rename(new_path)
        logger.info("Session directory renamed → %s", new_path)
        return new_path
    except OSError as err:
        logger.warning("Could not rename session dir: %s", err)
        return session_dir

# -----------------------------------------------------------------------
# Jira integration
# -----------------------------------------------------------------------


def push_to_jira(
    config: "JiraConfig",  # type: ignore[name-defined]
    markdown_content: str,
    clicks: list,
    mode: ProcessMode,
) -> list[str]:
    """Parse a generated Markdown report and push each task/step to Jira.

    Purpose:
        For QA mode, extracts each ``### Task N:`` block from the Markdown and
        creates one Jira issue per task.  For Documentation mode, creates a
        single Jira issue with the entire guide as its description.

        Screenshot attachments are matched by scanning each block for
        ``click_NNNN.png`` references.

    Args:
        config: Populated ``JiraConfig`` from ``audit_tool.config``.
        markdown_content: The full AI-generated Markdown string.
        clicks: Click records (used to resolve screenshot paths).
        mode: Process mode — determines parsing strategy.

    Returns:
        List of created Jira issue keys, in order.

    Side Effects:
        HTTP calls to the Jira instance (delegated to ``jira_client``).

    Error Behavior:
        Errors from ``jira_client`` are propagated to the caller in
        ``generate_report``, which catches and logs them without crashing.

    Determinism: Nondeterministic.
    Idempotency: No.
    Thread Safety: Yes.
    """
    from audit_tool.jira_client import JiraIssuePayload, push_session_to_jira

    # Build a screenshot path lookup: filename → Path
    screenshot_lookup: dict[str, Path] = {
        click.screenshot_path.name: click.screenshot_path
        for click in clicks
        if click.screenshot_path.exists()
    }

    payloads: list[JiraIssuePayload] = []

    if mode == ProcessMode.DOCUMENTATION:
        # Single issue: the whole guide
        payloads.append(JiraIssuePayload(
            summary="Documentation: " + _extract_doc_title(markdown_content),
            description_markdown=markdown_content,
            labels=["vibecheck", "documentation"],
            priority="Medium",
            attachments=list(screenshot_lookup.values()),
            task_number=0,
        ))
    else:
        # QA mode: one issue per ### Task N block
        payloads = _parse_qa_tasks_to_payloads(markdown_content, screenshot_lookup)

    return push_session_to_jira(config, payloads)


def _extract_doc_title(markdown_text: str) -> str:
    """Extract the H1 title from a Markdown string, or return a fallback."""
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped.lstrip("# ").strip()[:200]
    return "Walkthrough"


def _parse_qa_tasks_to_payloads(
    markdown_text: str,
    screenshot_lookup: dict[str, Path],
) -> list["JiraIssuePayload"]:
    """Split a QA-mode Markdown report into per-task ``JiraIssuePayload`` objects.

    Args:
        markdown_text: Full AI-generated Markdown.
        screenshot_lookup: Filename → Path for all captured screenshots.

    Returns:
        One ``JiraIssuePayload`` per detected ``### Task N`` block.
    """
    from audit_tool.jira_client import JiraIssuePayload

    # Split on ### Task headers
    task_pattern = re.compile(r"^### Task \d+[:\s]", re.MULTILINE)
    splits = task_pattern.split(markdown_text)
    headers = task_pattern.findall(markdown_text)

    payloads: list[JiraIssuePayload] = []

    for task_number, (header, body) in enumerate(zip(headers, splits[1:]), start=1):
        raw_title = header.strip().removeprefix("### Task ").strip()
        # Remove leading "1: " or "1 " numeric prefix
        raw_title = re.sub(r"^\d+[:\s]+", "", raw_title).strip()
        summary = f"[QA] {raw_title}" if raw_title else f"[QA] Task {task_number}"

        # Detect priority from body
        priority = "Medium"
        priority_match = re.search(
            r"Priority:\s*(Critical|High|Medium|Low)", body, re.IGNORECASE
        )
        if priority_match:
            raw_priority = priority_match.group(1).capitalize()
            # Jira doesn't have "Critical" by default; map to "Highest"
            priority = "Highest" if raw_priority == "Critical" else raw_priority

        # Find screenshot references in this block
        attachments: list[Path] = []
        for fname, fpath in screenshot_lookup.items():
            if fname in body:
                attachments.append(fpath)

        payloads.append(JiraIssuePayload(
            summary=summary[:255],
            description_markdown=body.strip(),
            labels=["vibecheck", "qa"],
            priority=priority,
            attachments=attachments,
            task_number=task_number,
        ))

    return payloads


# Needed for type references when ClickRecord hasn't been imported yet
from audit_tool.mouse_tracker import ClickRecord  # noqa: E402
from audit_tool.config import JiraConfig  # noqa: E402
